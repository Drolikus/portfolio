from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables"
OUT.mkdir(exist_ok=True)
FINAL = OUT / "Vladyslav_Kikhtenko_Portfolio_Brief.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9E2EF"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        borders.append(tag)
    tc_pr.append(borders)


def add_run(paragraph, text, bold=False, color="111827", size=10):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def style_paragraph(paragraph, before=0, after=4, line=1.05):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def heading(doc, text):
    p = doc.add_paragraph()
    style_paragraph(p, before=10, after=4)
    add_run(p, text.upper(), bold=True, color="0891B2", size=9)
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.62)
section.bottom_margin = Inches(0.62)
section.left_margin = Inches(0.68)
section.right_margin = Inches(0.68)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(9.6)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
style_paragraph(title, after=2)
add_run(title, "Vladyslav Kikhtenko", bold=True, color="0F172A", size=24)

subtitle = doc.add_paragraph()
style_paragraph(subtitle, after=8)
add_run(subtitle, "Junior Frontend Developer in Germany | HTML, CSS, JavaScript | Voltage guitar store", color="475569", size=10.5)

lead = doc.add_paragraph()
style_paragraph(lead, after=8, line=1.12)
add_run(
    lead,
    "I build polished responsive interfaces with visible proof: real project states, quality checks, command navigation, and honest scope. My main product build is Voltage, an online guitar store with catalog, cart, admin, and multilingual UI direction.",
    color="1F2937",
    size=9.8,
)

summary = doc.add_table(rows=1, cols=4)
summary.alignment = WD_TABLE_ALIGNMENT.CENTER
summary.autofit = False
labels = [
    ("12", "command actions"),
    ("9", "proof blocks"),
    ("7/7", "live QA audit"),
    ("4", "visitor routes/views"),
]
for idx, (value, label) in enumerate(labels):
    cell = summary.cell(0, idx)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "EEF9FC" if idx != 2 else "ECFDF3")
    set_cell_border(cell, "B7DDE8" if idx != 2 else "A7F3D0")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, after=0)
    add_run(p, value + "\n", bold=True, color="0891B2" if idx != 2 else "15803D", size=15)
    add_run(p, label, color="334155", size=8)

heading(doc, "Project proof")
projects = doc.add_table(rows=3, cols=3)
projects.alignment = WD_TABLE_ALIGNMENT.CENTER
projects.autofit = False
headers = ["Surface", "What it proves", "Current status"]
for i, text in enumerate(headers):
    cell = projects.cell(0, i)
    set_cell_shading(cell, "0F172A")
    set_cell_border(cell, "0F172A")
    p = cell.paragraphs[0]
    add_run(p, text, bold=True, color="F8FAFC", size=8.6)
rows = [
    ("Voltage", "E-commerce UI pressure: catalog, cart, admin flow, i18n labels.", "Active rebuild"),
    ("Portfolio", "Command palette, proof mode, visitor routes, contact builder, QA gate.", "Live proof surface"),
]
for r, row in enumerate(rows, start=1):
    for c, text in enumerate(row):
        cell = projects.cell(r, c)
        set_cell_border(cell)
        if c == 0:
            set_cell_shading(cell, "F8FAFC")
        p = cell.paragraphs[0]
        style_paragraph(p, after=0, line=1.08)
        add_run(p, text, bold=(c == 0), color="111827" if c == 0 else "334155", size=8.8)

heading(doc, "How I work")
work = doc.add_paragraph()
style_paragraph(work, after=4, line=1.13)
for item in [
    "Ship visible work, then tighten weak parts.",
    "Tie claims to proof instead of vague self-praise.",
    "Test mobile, overflow, assets, navigation, and interaction wiring.",
    "Keep future AI/Sentry features behind safe backend or configuration boundaries.",
]:
    work.add_run("• ").font.name = "Arial"
    add_run(work, item + "\n", color="334155", size=9.2)

heading(doc, "Best fit")
fit = doc.add_paragraph()
style_paragraph(fit, after=6, line=1.13)
add_run(fit, "Junior frontend role, UI implementation task, portfolio/code review, or practical frontend mentorship. Strongest current evidence: responsive static UI, vanilla JavaScript interactions, product storytelling, and quality guardrails.", color="334155", size=9.2)

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
style_paragraph(footer, before=8, after=0)
add_run(footer, "Portfolio is actively improving | Voltage is actively rebuilding | Contact route lives in the website", color="64748B", size=8)

doc.save(FINAL)
print(FINAL)
